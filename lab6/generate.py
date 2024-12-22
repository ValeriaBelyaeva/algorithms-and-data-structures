#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Пример кода на Python, который с помощью библиотеки python-docx
формирует docx-файл со структурой отчёта по задачам в соответствии
с описанным шаблоном.
"""

from docx import Document
from docx.shared import Pt

def generate_lorem_ipsum(word_count):
    """
    Возвращает строку из word_count слов (условный 'lorem ipsum').
    """
    base_words = (
        "lorem ipsum dolor sit amet consectetur adipiscing elit sed do eiusmod tempor "
        "incididunt ut labore et dolore magna aliqua ullamco laboris nisi ut aliquip ex ea "
        "commodo consequat duis aute irure dolor in reprehenderit in voluptate velit esse "
        "cillum dolore eu fugiat nulla pariatur excepteur sint occaecat cupidatat non proident "
        "sunt in culpa qui officia deserunt mollit anim id est laborum"
    ).split()
    # Если нужно больше слов, будем дублировать
    repeated = []
    while len(repeated) < word_count:
        repeated.extend(base_words)
    return " ".join(repeated[:word_count])

def generate_report_for_task(document, task_number, task_name_ru):
    """
    Добавляет в документ docx отчёт по одной задаче, согласно описанному паттерну.
    """

    # 1) Заголовок задачи
    header = f"Задача № {task_number}. {task_name_ru}"
    document.add_heading(header, level=1)

    # 2) Блок с "Функции из кода решения задачи..."
    #    Для демонстрации добавляем пример кода с псевдо-подсветкой.
    code_text = (
        "def function_example(param1, param2):\n"
        "    \"\"\"Пример функции\"\"\"\n"
        "    result = param1 + param2\n"
        "    return result\n"
    )
    code_paragraph = document.add_paragraph()
    code_run = code_paragraph.add_run(code_text)
    # Немного стилевого оформления (не полноценная подсветка, но хоть шрифт фиксированной ширины)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(10)

    # 3) Описание кода на 120 слов
    #    Генерируем условный текст lorem ipsum
    desc_text = "Описание кода: " + generate_lorem_ipsum(120)
    document.add_paragraph(desc_text)

    # 4) Таблица 1 (входные/выходные данные)
    table1 = document.add_table(rows=4, cols=3)
    table1.style = 'Light List Accent 1'
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = "тест"
    hdr_cells[1].text = "input.txt"
    hdr_cells[2].text = "output.txt"

    row_min = table1.rows[1].cells
    row_min[0].text = "минимальные вх"
    row_min[1].text = "*минимальные вх*"
    row_min[2].text = "*результат кода на минимальных вх*"

    row_example = table1.rows[2].cells
    row_example[0].text = "пример"
    row_example[1].text = "*вх примера из условия этой задачи*"
    row_example[2].text = "*результат кода на вх примера*"

    row_max = table1.rows[3].cells
    row_max[0].text = "максимальные вх"
    row_max[1].text = "*максимальные вх*"
    row_max[2].text = "*результат кода на максимальных вх*"

    document.add_paragraph()  # Пустая строка для разделения

    # 5) Таблица 2 (затраты памяти / времени)
    table2 = document.add_table(rows=4, cols=3)
    table2.style = 'Light List Accent 2'
    hdr2_cells = table2.rows[0].cells
    hdr2_cells[0].text = "тест"
    hdr2_cells[1].text = "Затраты памяти (Мб)"
    hdr2_cells[2].text = "Время выполнения (с)"

    row_lower = table2.rows[1].cells
    row_lower[0].text = "Нижняя граница"
    row_lower[1].text = "*соответствующие данные*"
    row_lower[2].text = "*соответствующие данные*"

    row_ex = table2.rows[2].cells
    row_ex[0].text = "пример"
    row_ex[1].text = "*соответствующие данные*"
    row_ex[2].text = "*соответствующие данные*"

    row_upper = table2.rows[3].cells
    row_upper[0].text = "Верхняя граница"
    row_upper[1].text = "*соответствующие данные*"
    row_upper[2].text = "*соответствующие данные*"

    document.add_paragraph()  # Пустая строка

    # 6) Вывод по задаче (60 слов)
    conclusion_text = "Вывод по задаче: " + generate_lorem_ipsum(60)
    document.add_paragraph(conclusion_text)

def main():
    document = Document()

    # Пример списка задач (номер, название)
    tasks = [
        (1, "Множество"),
        (2, "Телефонная книга"),
        (4, "Прошитый ассоциативный массив"),
        (7, "Драгоценные камни"),
        (8, "Почти интерактивная хеш-таблица")
    ]

    for tnum, tname in tasks:
        generate_report_for_task(document, tnum, tname)

    # Сохраняем итоговый документ
    document.save("report.docx")

if __name__ == "__main__":
    main()
